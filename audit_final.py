"""Final audit of Paper D v3 manuscript against pipeline data."""
import pandas as pd
import numpy as np
from scipy import stats
from docx import Document
import re
import os

print("=" * 70)
print("PAPER D v3 FINAL AUDIT (314 sites)")
print("=" * 70)

df = pd.read_csv('outputs/toa_tables/site_summary_v3_full.csv')
doc = Document('paper_d/manuscript/Shahid_2026_Surface_TOA_Transfer_v3_FINAL.docx')
all_text = ' '.join([p.text for p in doc.paragraphs])

alpha = 'forcing_proxy'
cre = 'ceres_toa_cre_net_mean'
cre_sw = 'ceres_toa_cre_sw_mean'

errors, warnings, passes = [], [], []

def check(label, ok, detail=""):
    (passes if ok else errors).append(f"{'PASS' if ok else 'FAIL'}: {label}" + (f" -- {detail}" if not ok else ""))

def warn(label, detail):
    warnings.append(f"WARN: {label} -- {detail}")

# 1. SITE COUNT
n = len(df)
check("Site count = 314", n == 314, f"got {n}")
check("'314' in manuscript", "314" in all_text)
check("No stale '313 FLUXNET'", "313 FLUXNET" not in all_text, "313 still in text")
check("No stale '255 FLUXNET'", "255 FLUXNET" not in all_text, "255 still in text")

# 2. BIOME COUNTS
bc = df['biome_group'].value_counts()
check(f"Forest = {bc.get('Forest',0)}", bc.get('Forest',0) == 132, f"got {bc.get('Forest')}")
check("'132 forest' in text", "132 forest" in all_text.lower(), "not found")

# 3. KEY REGRESSIONS
mask = df[alpha].notna() & df[cre].notna()
x, y = df.loc[mask, alpha].values, df.loc[mask, cre].values
s, i, r, p, se = stats.linregress(x, y)
check(f"CRE_net gamma ~ 8.40 (got {s:.2f})", abs(s - 8.40) < 0.15)
check(f"CRE_net R2 ~ 0.130 (got {r**2:.3f})", abs(r**2 - 0.130) < 0.01)
check(f"CRE_net p ~ 4.9e-11 (got {p:.1e})", p < 1e-10)

mask2 = df[alpha].notna() & df[cre_sw].notna()
s2, _, r2, p2, _ = stats.linregress(df.loc[mask2, alpha], df.loc[mask2, cre_sw])
check(f"CRE_SW gamma ~ 11.80 (got {s2:.2f})", abs(s2 - 11.80) < 0.15)
check(f"CRE_SW R2 ~ 0.176 (got {r2**2:.3f})", abs(r2**2 - 0.176) < 0.01)

# 4. SEASONAL
check("JJA gamma '18.3' in text", "18.3" in all_text, "not found")
check("DJF not significant in text", "0.71" in all_text or "0.70" in all_text, "DJF p not found")
check("'14 times' in text", "14 times" in all_text, "not found")

# 5. ERA5 MEDIATION
check("'54%' mediation in text", "54%" in all_text, "not found")
check("'30%' BLH in text", "30%" in all_text, "not found")
check("'8.4e-5' or similar in text", "8.4e-5" in all_text or "8.4e-05" in all_text, "not found")

# 6. HOMOGENEITY
check("R2=0.633 in text", "0.633" in all_text, "not found")
check("n=71 in text", "n = 71" in all_text or "n=71" in all_text, "not found")

# 7. FINANCIALS
f_alpha = df.loc[df['biome_group']=='Forest', alpha].mean()
s_alpha = df.loc[df['biome_group']=='Shrubland', alpha].mean()
check(f"Forest alpha ~ 0.856 (got {f_alpha:.3f})", abs(f_alpha - 0.856) < 0.02)
check("Financial $83 or $84 in text", "$83" in all_text or "$84" in all_text, "not found")

# 8. BR-SA1
check("'BR-Sa1' in text", "BR-Sa1" in all_text, "not found")
check("'most' + 'studied' in text", "most" in all_text.lower() and "studied" in all_text.lower(), "not found")
check("'Km67' or 'Santarem' in text", "Km67" in all_text or "Santarem" in all_text, "not found")

# 9. CLEAR-SKY NOT SIGNIFICANT
mask_c = df[alpha].notna() & df['ceres_toa_net_clr_mean'].notna()
_, _, _, p_clr, _ = stats.linregress(df.loc[mask_c, alpha], df.loc[mask_c, 'ceres_toa_net_clr_mean'])
check(f"Clear-sky TOA not significant (p={p_clr:.3f})", p_clr > 0.05)

# 10. HEADER/DATE
header_text = ''.join(p.text for s in doc.sections for p in s.header.paragraphs)
check("No 'v2' in header", 'v2' not in header_text, f"header: '{header_text}'")
check("No '(v3)' in date", '(v3)' not in all_text[:500], "version tag in date")

# 11. IMAGES
img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
check(f"Images >= 9 (got {img_count})", img_count >= 9)

# 12. REFERENCES
check("Cerasoli cited", "Cerasoli" in all_text)
check("Duveiller 2021 cited", "Duveiller" in all_text)
check("Teuling cited", "Teuling" in all_text)
check("Blichner cited", "Blichner" in all_text)

# 13. DATA AVAILABILITY
check("Zenodo DOI", "10.5281/zenodo.19328341" in all_text)
check("FLUXNET DOI", "10.5281/zenodo.10885933" in all_text)
check("ORCID", "0009-0003-9709-4241" in all_text)
check("AI disclosure (Claude/Anthropic)", "Claude" in all_text and "Anthropic" in all_text)

# 14. CODE QUALITY
if os.path.exists('data/processed/synthetic_sites_with_toa.csv'):
    warn("Synthetic data file exists", "data/processed/synthetic_sites_with_toa.csv should be deleted before git")

for src in ['src/toa_transfer.py', 'src/robustness.py', 'src/figures_v3.py', 'run_toa_analysis.py']:
    if os.path.exists(src):
        with open(src) as f:
            content = f.read()
        if 'chagpt' in content.lower():
            warn(f"Hardcoded path in {src}", "Contains project-specific paths")

# REPORT
print()
print(f"PASSES: {len(passes)}")
for p in passes:
    print(f"  [{p}]")

if warnings:
    print(f"\nWARNINGS: {len(warnings)}")
    for w in warnings:
        print(f"  [{w}]")

if errors:
    print(f"\nERRORS: {len(errors)}")
    for e in errors:
        print(f"  [{e}]")
else:
    print("\nERRORS: 0")

print()
print("=" * 70)
verdict = "READY FOR SUBMISSION" if len(errors) == 0 else f"FIX {len(errors)} ERRORS BEFORE SUBMISSION"
print(f"VERDICT: {verdict}")
print("=" * 70)
